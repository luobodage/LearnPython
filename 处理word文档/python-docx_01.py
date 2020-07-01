import os
from docx import Document

#要读取的文件的上层文件夹路径
fileDir = "C:\\Users\\萝卜ovo\\Desktop\\杂\\赵如赵老师\\Test"

# 获取文件地址
def GetFileList(dir, fileList):
    newDir = dir
    if os.path.isfile(dir):
        fileList.append(dir)
    elif os.path.isdir(dir):
        for s in os.listdir(dir):
            # 如果需要忽略某些文件夹，使用以下代码
            # if s == "xxx":
            # continue
            newDir = os.path.join(dir, s)
            GetFileList(newDir, fileList)
    return fileList

#按行输出word内容
def printdocumet():
    for para in document.paragraphs:
        print(para.text)


#转换docx并且保存+test
def Convertdocument():
    #遍历word文档 并且替换关键词
    for x in document.paragraphs:
        x.text = x.text.replace('1、', '1.')
        x.text = x.text.replace('2、', '2.')
        x.text = x.text.replace('3、', '3.')
        x.text = x.text.replace('4、', '4.')
        x.text = x.text.replace('5、', '5.')
        x.text = x.text.replace('6、', '6.')
        x.text = x.text.replace('7、', '7.')
        x.text = x.text.replace('8、', '8.')
        x.text = x.text.replace('9、', '9.')
        x.text = x.text.replace('0、', '0.')
        x.text = x.text.replace('（', '(')
        x.text = x.text.replace('）', ').')
        x.text = x.text.replace('√', '对')
        x.text = x.text.replace('╳', '错')
        x.text = x.text.replace('答：', '[参考答案]')
        x.text = x.text.replace('一、填空', '[填空题]')
        x.text = x.text.replace('二、单项选择', '[单选题]')
        x.text = x.text.replace('三、多项选择', '[多选题]')
        x.text = x.text.replace('四、判断', '[判断题]')
        x.text = x.text.replace('五、简答题', '[简答题]')
        x.text = x.text.replace('六、应用题', '[简答题]')
    document.save(newdocxpath)

if __name__ == '__main__':
    # printdocumet()
    list = GetFileList(fileDir, [])
    #获取文件路径filepath 文件名filename 后缀名extension
    (filepath, tempfilename) = os.path.split(list[0])
    (filename, extension) = os.path.splitext(tempfilename)
    #获得文件路径docxpath1 newdocxpath1
    docxpath_info = filepath + "/" + filename + extension
    newdocxpath_info = filepath + "/" + filename + "test" + extension
    #整理成document能识别的格式
    docxpath = str(docxpath_info.replace("\\", "/"))
    newdocxpath = str(newdocxpath_info.replace("\\", "/"))
    document = Document(docxpath)
    #进行转换并且保存
    Convertdocument()

