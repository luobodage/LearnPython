from docx import Document


# 设置行和列的数量

rowNum = 2

colNum = 2

# 创建空白docx文档

doc = Document()

# 添加新段落

p = doc.add_paragraph('')


# 添加表格

table = doc.add_table(rows=rowNum, cols=colNum , style='Table Grid')

# 为表格的每个单元格添加文本

for row in range(rowNum):

    for col in range(colNum):

        cell = table.cell(row, col)

        cell.text = 'ok'

# 保存文件

doc.save('test1.docx')
